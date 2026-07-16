import sys
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    
    cell = table.cell(0, 0)
    set_cell_background(cell, "F1F5F9")  # Light gray
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    # Border styling (thin left border or subtle box)
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="3B82F6"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    
    # Add an empty paragraph after table for spacing
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(4)
    p_after.paragraph_format.space_after = Pt(6)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)  # Deep Slate Navy
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)  # Amber / Golden accent
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)  # Blue accent
    return p

def add_body_p(doc, text, bold_prefix=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.18
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = 'Arial'
        r_bold.font.size = Pt(10.5)
        r_bold.font.bold = True
        r_bold.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    return p

def add_bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = 'Arial'
        r_bold.font.size = Pt(10.5)
        r_bold.font.bold = True
        r_bold.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    return p

def main():
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # --- TITLE & HEADER ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(10)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("TÀI LIỆU HƯỚNG DẪN TÍCH HỢP FRONTEND & API SPECIFICATION")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(24)
    run_sub = sub_p.add_run("Dự án: Website B2B Xuất Khẩu Thủy Hải Sản Golden Seafood Co., Ltd\nĐối tượng: Frontend Engineer / UI/UX Lead | Phiên bản API: v1.0.0")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # --- SECTION 1 ---
    add_heading_1(doc, "1. TỔNG QUAN HỆ THỐNG & CHIẾN LƯỢC DATA")
    add_body_p(doc, "Hệ thống Golden Seafood được thiết kế tối ưu hóa chuyên biệt cho mô hình B2B Global Export (Xuất khẩu toàn cầu). Mục tiêu cốt lõi là đạt tốc độ tải trang cực nhanh (< 2 giây) và đem đến trải nghiệm sang trọng, uy tín vượt bậc cho khách mua hàng quốc tế.")
    
    add_heading_2(doc, "1.1 Phân định rõ ràng Dữ liệu Động (API) vs Dữ liệu Tĩnh (Hardcode FE)")
    add_body_p(doc, "Để đảm bảo hiệu năng tải trang dưới 100ms theo đúng tinh thần tối ưu đã thống nhất, hệ thống phân chia rõ ràng:")
    
    add_bullet(doc, " Trang Chủ (Home): Khối giới thiệu Về chúng tôi, quy trình làm việc 5 bước, lý do chọn Golden Seafood.", "Dữ Liệu Tĩnh (Hardcode FE / JSON):")
    add_bullet(doc, " Trang Dịch vụ Gia công (Processing Services): Quy trình chế biến sâu cá thịt trắng (Whitefish Deep-Processing), đóng gói chuẩn siêu thị quốc tế (Retail-ready).", "", level=1)
    add_bullet(doc, " Trang Đảm bảo Chất lượng (Quality Assurance): Quy trình kiểm định QC 4 bước, tiêu chuẩn nhà máy (HACCP, BRC, ASC, BAP).", "", level=1)
    add_bullet(doc, " Trang Liên hệ (Contact Us) & Khối Call-out Mixed Container Footer ở cuối trang chi tiết sản phẩm.", "", level=1)

    add_bullet(doc, " Danh mục sản phẩm (Categories): 5 nhóm ngành hàng xuất khẩu cốt lõi kèm số lượng sản phẩm đang hoạt động.", "Dữ Liệu Động (Gọi từ Backend API):")
    add_bullet(doc, " Danh sách & Chi tiết Sản phẩm (Products & Specs): Tên song ngữ Anh/Việt, phân loại raw / cooked / value_added, hình ảnh gallery và bảng thông số kỹ thuật chi tiết.", "", level=1)
    add_bullet(doc, " Hệ thống Báo giá B2B (Inquiry Basket & Contact Form): Gửi form báo giá giỏ hàng chứa nhiều sản phẩm cùng lúc, upload file spec sheet từ buyer và nhận auto-reply email.", "", level=1)

    add_heading_2(doc, "1.2 Cấu hình Môi trường Frontend (.env.local)")
    create_code_block(doc, """# Môi trường Development
NEXT_PUBLIC_API_BASE_URL=http://localhost:3000/api
NEXT_PUBLIC_UPLOADS_URL=http://localhost:3000/

# Môi trường Production
# NEXT_PUBLIC_API_BASE_URL=https://api.goldenseafood.com.vn/api""")

    # --- SECTION 2 ---
    add_heading_1(doc, "2. TỔNG HỢP TOÀN BỘ PUBLIC API ENDPOINTS")
    add_body_p(doc, "Mọi request từ Frontend gửi lên API (ngoại trừ upload file multipart/form-data) đều dùng HTTP Header: Content-Type: application/json.")

    # Table 1: Summary of APIs
    add_heading_2(doc, "2.1 Bảng Tổng Hợp Nhanh Các Endpoints")
    
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    col_widths = [Inches(1.1), Inches(1.8), Inches(1.6), Inches(2.0)]
    for row in table.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = w

    headers = ["Nhóm API", "Phương thức & Endpoint", "Tham số / Body JSON", "Mô tả & Cách dùng cho FE"]
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "0F172A")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    api_rows = [
        ("Danh Mục", "GET /api/categories", "Không có", "Trả về 5 danh mục chính kèm productCount. Dùng cho Sidebar lọc trang Sản phẩm."),
        ("Sản Phẩm", "GET /api/products", "Query Params:\n• page=1&limit=9\n• category_id=1\n• type=raw/cooked\n• search=từ khóa", "Lọc sản phẩm AJAX không cần tải lại trang. Trả về mảng products và pagination."),
        ("Sản Phẩm", "GET /api/products/featured", "Không có", "Trả về danh sách sản phẩm nổi bật (is_featured: true) cho section Trang Chủ."),
        ("Chi Tiết SP", "GET /api/products/:slug", "URL Param:\n• :slug", "Trả về chi tiết 1 sản phẩm kèm gallery ảnh images và bảng specifications."),
        ("Liên Hệ", "POST /api/inquiries/contact", "Body JSON:\n{ full_name, company_name, email, country, whatsapp_number, message }", "Form liên hệ nhanh trang Contact Us (giới hạn 10 request/15 phút chống spam)."),
        ("Giỏ Báo Giá", "POST /api/inquiries/basket", "Body JSON:\n{ full_name, company_name, email, country, whatsapp_number, destination_port, special_requirements, attachment_url, items: [...] }", "API cốt lõi B2B: Gửi danh sách N mặt hàng trong giỏ báo giá. Tự động lưu DB, sinh mã GS-YYYYMMDD-XXX và gửi email kép."),
        ("Upload File", "POST /api/upload", "Form-data: file\n(Max 10MB: PDF, Excel, Word, JPG)", "Upload spec sheet / yêu cầu đóng gói của buyer. FE gọi lấy data.url trước khi submit form báo giá.")
    ]

    for row_idx, data in enumerate(api_rows):
        row_cells = table.add_row().cells
        bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for i, text in enumerate(data):
            row_cells[i].text = text
            row_cells[i].width = col_widths[i]
            set_cell_background(row_cells[i], bg_color)
            set_cell_margins(row_cells[i], top=100, bottom=100, left=100, right=100)
            p = row_cells[i].paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(9.0)
                run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            if i == 1: # Highlight endpoint name
                p.runs[0].font.name = 'Consolas'
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    # Style table borders
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(f'''
                <w:tcBorders {nsdecls("w")}>
                    <w:top w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>
                    <w:left w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>
                    <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>
                    <w:right w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>
                </w:tcBorders>
            ''')
            tcPr.append(tcBorders)

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(6)

    # 2.2 Detailed Schemas
    add_heading_2(doc, "2.2 Chi Tiết Schema Request & Response Cốt Lõi")
    
    add_heading_3(doc, "GET /api/products (Danh sách & Lọc sản phẩm)")
    create_code_block(doc, """// Request: GET /api/products?category_id=1&type=value_added&page=1&limit=9
// Response (200 OK):
{
  "success": true,
  "data": {
    "products": [
      {
        "id": 1,
        "category_id": 1,
        "name_en": "Frozen Vannamei Shrimp - Breaded",
        "name_vi": "Tôm thẻ chân trắng tẩm bột đông lạnh",
        "slug": "frozen-vannamei-shrimp-breaded",
        "short_desc_en": "Premium IQF breaded Vannamei shrimp, crisp golden coating...",
        "thumbnail_url": "https://images.unsplash.com/photo-1565680018434-b513d5e5fd47?w=800",
        "product_type": "value_added",
        "is_featured": true,
        "category": { "id": 1, "name_en": "Shrimp", "name_vi": "Tôm" }
      }
    ],
    "pagination": { "page": 1, "limit": 9, "total": 12, "totalPages": 2 }
  }
}""")

    add_heading_3(doc, "GET /api/products/:slug (Chi tiết sản phẩm & Bảng thông số)")
    create_code_block(doc, """// Request: GET /api/products/frozen-vannamei-shrimp-breaded
// Response (200 OK):
{
  "success": true,
  "data": {
    "id": 1,
    "name_en": "Frozen Vannamei Shrimp - Breaded",
    "name_vi": "Tôm thẻ chân trắng tẩm bột đông lạnh",
    "slug": "frozen-vannamei-shrimp-breaded",
    "description_en": "Sourced from ASC/BAP certified farms in Vietnam...",
    "thumbnail_url": "https://images.unsplash.com/...",
    "product_type": "value_added",
    "images": [
      { "id": 1, "image_url": "https://images.unsplash.com/...", "is_primary": true }
    ],
    "specifications": [
      { "spec_key_en": "Scientific Name", "spec_key_vi": "Tên khoa học", "spec_value": "Litopenaeus vannamei" },
      { "spec_key_en": "Origin", "spec_key_vi": "Xuất xứ", "spec_value": "Vietnam (Farm-raised)" },
      { "spec_key_en": "Size / Counting", "spec_key_vi": "Kích cỡ", "spec_value": "16/20, 21/25, 26/30, 31/40" },
      { "spec_key_en": "Freezing Method", "spec_key_vi": "Quy cách cấp đông", "spec_value": "IQF / Block Frozen" }
    ]
  }
}""")

    add_heading_3(doc, "POST /api/inquiries/basket (Gửi Giỏ Hàng Báo Giá B2B)")
    create_code_block(doc, """// Request Body (JSON):
{
  "full_name": "Elena Rostova",
  "job_title": "Procurement Manager",
  "company_name": "Vostok Seafood Trading",
  "email": "elena@vostokseafood.ru",
  "country": "Russia",
  "whatsapp_number": "+7-999-123-4567",
  "destination_port": "St. Petersburg Port",
  "message": "Need mixed container quote. Please include target freight rates.",
  "special_requirements": "All boxes must have Russian language labels and EAC marks.",
  "attachment_url": "uploads/specs/spec-sheet-1784176444.pdf",
  "items": [
    {
      "product_id": 1,
      "product_name": "Frozen Vannamei Shrimp - Breaded",
      "quantity": 25,
      "specifications": "Size 16/20, IQF, Glazing 10%, Bulk carton 10kg",
      "notes": "Need ASC certificate copy attached with Proforma Invoice"
    },
    {
      "product_id": 4,
      "product_name": "Whole Cleaned Squid - Frozen",
      "quantity": 15,
      "specifications": "Size 8/12 cm, IQF, Glazing 5%, Block 7.5kg"
    }
  ]
}

// Response (201 Created):
{
  "success": true,
  "message": "Inquiry basket submitted successfully. We will get back to you within 24 working hours.",
  "data": {
    "id": 2,
    "inquiry_code": "GS-20260716-002",
    "status": "new",
    "itemCount": 2
  }
}""")

    # --- SECTION 3 ---
    add_heading_1(doc, "3. TYPESCRIPT INTERFACES CHUẨN CHO FRONTEND")
    add_body_p(doc, "FE Developer có thể copy trực tiếp trọn bộ Interfaces này vào file types/api.ts trong dự án Next.js / React để đảm bảo Type-Safety 100%:")
    
    create_code_block(doc, """export interface IApiResponse<T> {
  success: boolean;
  message?: string;
  data: T;
}

export interface IPagination {
  page: number; limit: number; total: number; totalPages: number;
}

export interface ICategory {
  id: number; name_en: string; name_vi: string; slug: string;
  icon_url: string | null; sort_order: number; is_active: boolean; productCount?: number;
}

export type ProductType = 'raw' | 'cooked' | 'value_added';

export interface IProductSpecification {
  id?: number; spec_key_en: string; spec_key_vi: string; spec_value: string; sort_order?: number;
}

export interface IProductImage {
  id?: number; image_url: string; is_primary: boolean; sort_order: number;
}

export interface IProduct {
  id: number; category_id: number; name_en: string; name_vi: string; slug: string;
  short_desc_en: string | null; short_desc_vi: string | null;
  description_en?: string | null; description_vi?: string | null;
  thumbnail_url: string | null; product_type: ProductType; is_featured: boolean;
  category?: ICategory; images?: IProductImage[]; specifications?: IProductSpecification[];
}

export interface IBasketItemState {
  product_id: number; product_name: string; product_slug: string;
  thumbnail_url: string | null; quantity?: number; specifications?: string; notes?: string;
}

export interface IInquiryBasketPayload {
  full_name: string; job_title?: string; company_name: string; email: string;
  country: string; whatsapp_number?: string; destination_port?: string;
  message?: string; special_requirements?: string; attachment_url?: string | null;
  items: { product_id: number; product_name?: string; quantity?: number; specifications?: string; notes?: string; }[];
}""")

    # --- SECTION 4 ---
    add_heading_1(doc, "4. HƯỚNG DẪN THIẾT KẾ & NGHIỆP VỤ UI/UX TỪNG TRANG (THEO EXCEL)")
    
    add_heading_2(doc, "4.1 Ngôn ngữ & Tông màu thương hiệu")
    add_bullet(doc, " Tông màu chủ đạo (Primary): Xanh đại dương sâu (#0B1D3A hoặc #0F172A) — Thể hiện uy tín hải sản đại dương.", "Màu sắc:")
    add_bullet(doc, " Tông màu điểm nhấn (Secondary/Accent): Vàng Champagne (#D97706 hoặc #F59E0B) — Thể hiện sự sang trọng vàng cao cấp của Golden Seafood.", "")
    add_bullet(doc, " Chế độ song ngữ (i18n): Nút Toggle EN / VI trên Navbar. Khi chọn Tiếng Anh render trường *_en, khi chọn Tiếng Việt render *_vi.", "Ngôn ngữ:")

    add_heading_2(doc, "4.2 Trang Sản Phẩm (Product Listing Page)")
    add_bullet(doc, " Bộ lọc thông minh bên trái (Sidebar Filter AJAX): Fetch 5 danh mục từ GET /api/categories hiển thị dạng list kèm badge số lượng. Khi click vào danh mục hoặc tick chọn Attribute Filter (Raw / Cooked / Value-added), gọi ngay GET /api/products?category_id=...&type=... để lọc tức thì mà không tải lại trang.", "Bộ lọc AJAX:")
    add_bullet(doc, " Hiệu ứng Hover B2B (Chuẩn Excel dòng 28): Khi rê chuột vào Thẻ Sản Phẩm (Product Card), ảnh phóng to nhẹ 5% (scale: 1.05), viền khung chuyển sang màu Vàng Champagne mỏng, và xuất hiện ngay nút hành động: + Add to Inquiry List (thay cho nút mua hàng thông thường).", "Hiệu ứng Hover:")

    add_heading_2(doc, "4.3 Trang Chi Tiết Sản Phẩm & Quy tắc Bảng Thông Số")
    add_body_p(doc, "Quy tắc hiển thị Bảng Thông Số Kỹ Thuật (Specification Table theo Excel dòng 42): Thiết kế bảng HTML sạch sẽ, tuyệt đối không dùng đường kẻ dọc đen đậm, chỉ dùng đường kẻ ngang mờ (border-bottom: 1px solid #E2E8F0), xen kẽ dòng trắng (#FFFFFF) và dòng xám siêu nhạt (#F8FAFC).")
    add_body_p(doc, "Ở cuối trang chi tiết sản phẩm, FE hardcode khối Call-out Container Hỗn Hợp (Mixed Container Footer):")
    create_code_block(doc, """<!-- Mixed Container Footer Banner (Hardcode FE) -->
<div class="bg-slate-900 text-white p-6 rounded-xl mt-8 border-l-4 border-amber-500 flex justify-between items-center">
  <div>
    <h4 class="text-lg font-bold text-amber-400">Looking for a flexible supply?</h4>
    <p class="text-slate-300 text-sm mt-1">We specialize in consolidating Mixed Containers (Shrimp, Fish, Cephalopods) to optimize your inventory and freight costs.</p>
  </div>
  <button onclick="openInquiryDrawer()" class="bg-amber-500 hover:bg-amber-600 text-slate-950 font-semibold px-5 py-2.5 rounded-lg">Request Mixed Quote →</button>
</div>""")

    add_heading_2(doc, "4.4 Widget Giỏ Hàng Báo Giá & Nút WhatsApp Instant Quote")
    add_body_p(doc, "Khi khách bấm gửi form báo giá POST /api/inquiries/basket thành công, hệ thống hiển thị Popup Chúc Mừng (Success Confirmation Modal) kèm Mã báo giá tự động GS-YYYYMMDD-XXX. Đồng thời, hiển thị ngay nút kết nối WhatsApp Instant Quote theo đúng yêu cầu Excel dòng 69:")
    create_code_block(doc, """<!-- Nút kết nối nhanh WhatsApp xuất hiện trong Success Popup -->
<a href="https://wa.me/84900000000?text=Hi%20Golden%20Seafood,%20I%20have%20just%20submitted%20inquiry%20basket%20code%20GS-20260716-002.%20Please%20provide%20a%20quick%20quote." target="_blank" class="btn-whatsapp-instant">
  💬 Chat directly via WhatsApp for instant quote →
</a>""")

    # --- SECTION 5 ---
    add_heading_1(doc, "5. HƯỚNG DẪN CLIENT STATE MANAGEMENT (ZUSTAND STORE)")
    add_body_p(doc, "Để giữ danh sách món hàng trong giỏ khi khách chuyển trang từ Tôm sang Mực, FE bắt buộc phải dùng Client State Management có lưu localStorage. Dưới đây là code mẫu Zustand store cho Next.js / React:")
    create_code_block(doc, """// store/useInquiryStore.ts
import { create } from 'zustand';
import { persist } from 'zustand/middleware';
import { IBasketItemState } from '@/types/api';

interface InquiryStore {
  items: IBasketItemState[]; isOpenDrawer: boolean;
  addItem: (item: IBasketItemState) => void;
  removeItem: (productId: number) => void;
  updateItemQuantity: (productId: number, quantity: number) => void;
  clearBasket: () => void;
  toggleDrawer: (open?: boolean) => void;
}

export const useInquiryStore = create<InquiryStore>()(
  persist(
    (set, get) => ({
      items: [], isOpenDrawer: false,
      addItem: (item) => {
        const current = get().items;
        const index = current.findIndex(i => i.product_id === item.product_id);
        if (index > -1) {
          const updated = [...current];
          updated[index].quantity = (updated[index].quantity || 1) + (item.quantity || 1);
          set({ items: updated, isOpenDrawer: true });
        } else {
          set({ items: [...current, { ...item, quantity: item.quantity || 1 }], isOpenDrawer: true });
        }
      },
      removeItem: (id) => set({ items: get().items.filter(i => i.product_id !== id) }),
      updateItemQuantity: (id, q) => set({ items: get().items.map(i => i.product_id === id ? { ...i, quantity: q } : i) }),
      clearBasket: () => set({ items: [] }),
      toggleDrawer: (open) => set({ isOpenDrawer: open !== undefined ? open : !get().isOpenDrawer })
    }),
    { name: 'golden-seafood-inquiry-basket' }
  )
);""")

    # --- SECTION 6 ---
    add_heading_1(doc, "6. KẾ HOẠCH & TIMELINE TRIỂN KHAI CHO FRONTEND (7 NGÀY)")
    add_body_p(doc, "Đội ngũ Frontend có thể bám sát checklist 5 giai đoạn dưới đây để hoàn thiện dự án đúng tiến độ và khớp 100% với Backend:")
    
    add_bullet(doc, " Khởi tạo project Next.js 14 App Router / Vite React + TailwindCSS, cấu hình font Inter & Outfit.", "Phase 1: Setup Project & Base Components (Ngày 1):")
    add_bullet(doc, " Viết tiện ích API Client (axios/fetch có error handling), thiết lập đa ngôn ngữ i18n Anh - Việt và xây dựng Navbar/Footer B2B.", "", level=1)

    add_bullet(doc, " Cài đặt Zustand store kết nối localStorage, làm Floating Icon giỏ hàng ở góc phải màn hình kèm badge counter.", "Phase 2: Floating Basket Drawer & State (Ngày 2):")
    add_bullet(doc, " Tạo Component Slide-over Drawer nhập thông tin Buyer B2B, upload file spec sheet và gửi POST /api/inquiries/basket.", "", level=1)

    add_bullet(doc, " Component Product Card với hiệu ứng Hover B2B phóng to 5% + viền vàng champagne + nút Add to Inquiry List.", "Phase 3: Trang Danh Mục & Danh Sách SP (Ngày 3 - 4):")
    add_bullet(doc, " Trang /products với Sidebar lọc AJAX theo danh mục và thuộc tính Raw/Cooked/Value-added + Thanh tìm kiếm song ngữ + Phân trang.", "", level=1)

    add_bullet(doc, " Trang /products/[slug], Image Gallery với Zoom-in khi rê chuột, render Bảng thông số kỹ thuật song ngữ xen kẽ xám/trắng.", "Phase 4: Trang Chi Tiết SP & Báo Giá (Ngày 5):")
    add_bullet(doc, " Nút Add to Inquiry Basket truyền dữ liệu vào store và hiển thị banner Mixed Container Footer ở cuối trang.", "", level=1)

    add_bullet(doc, " Hardcode song ngữ các trang tĩnh (Trang Chủ, Gia Công, Chất Lượng, Liên Hệ), kiểm thử Responsive trên mọi thiết bị.", "Phase 5: Các Trang Tĩnh & Polish Core Web Vitals (Ngày 6 - 7):")
    add_bullet(doc, " Đo điểm Lighthouse Performance > 95 điểm, đảm bảo tốc độ tải trang dưới 2 giây đúng chỉ tiêu Excel.", "", level=1)

    # Output paths
    output_path = os.path.join(os.path.abspath("."), "GoldenSeafood_TaiLieu_TichHop_Frontend.docx")
    doc.save(output_path)
    print(f"✅ Created Word document successfully at: {output_path}")

if __name__ == "__main__":
    main()
